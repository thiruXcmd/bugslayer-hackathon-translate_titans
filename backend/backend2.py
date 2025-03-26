from flask import Flask, request, jsonify
import os
from docx import Document
from deep_translator import GoogleTranslator

app = Flask(__name__)
UPLOAD_FOLDER = './uploads'
TRANSLATED_FOLDER = './processed'


os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(TRANSLATED_FOLDER, exist_ok=True)

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400

    
    upload_path = os.path.join(UPLOAD_FOLDER, file.filename)
    file.save(upload_path)

    
    translated_path = os.path.join(TRANSLATED_FOLDER, f"processed_{file.filename}")
    INDIAN_LANGUAGES = {
        "Hindi": "hi",
        "Bengali": "bn",
        "Telugu": "te",
        "Marathi": "mr",
        "Tamil": "ta",
        "Gujarati": "gu",
        "Kannada": "kn",
        "Malayalam": "ml",
        "Punjabi": "pa",
        "Urdu": "ur",
        "Odia": "or",
        "Assamese": "as",
        "Maithili": "mai"
    }

    def allowed_file(filename):
        """Check if the uploaded file has a valid extension."""
        return filename.lower().endswith(".docx")

    def read_word_file(filepath):
        """Read a .docx file and extract text."""
        doc = Document(filepath)
        return "\n".join([para.text for para in doc.paragraphs])

    def write_word_file(filepath, text):
        """Write translated text to a new .docx file."""
        doc = Document()
        doc.add_paragraph(text)
        doc.save(filepath)

    def translate_text(text, target_lang):
        """Translate text using Google Translate API."""
        translator = GoogleTranslator(source="auto", target=target_lang)
        return translator.translate(text)

    def translate_and_save(file_path, target_language):
        """Read, translate, and save a document."""
        if not allowed_file(file_path):
            print(" Error: Unsupported file format. Please use .docx")
            return None

        print(f" Reading file: {file_path}")
        original_text = read_word_file(file_path)

        if not original_text.strip():
            print(" Error: No content to translate!")
            return None

        print(" Translating text...")
        translated_text = translate_text(original_text, target_language)

        translated_filename = f"translated_{target_language}.docx"
        translated_filepath = os.path.join(TRANSLATED_FOLDER, translated_filename)

        print(f" Saving translated file: {translated_filepath}")
        write_word_file(translated_filepath, translated_text)

        print(f" Translation complete! Saved as {translated_filepath}")
        return translated_filepath

   
    print("\n Available Indian languages for translation:")
    for i, (lang, code) in enumerate(INDIAN_LANGUAGES.items(), start=1):
        print(f"{i}. {lang}")

    
    while True:
        try:
            choice = int(input("\nEnter the number of the language you want to translate to: "))
            if 1 <= choice <= len(INDIAN_LANGUAGES):
                target_language = list(INDIAN_LANGUAGES.values())[choice - 1]
                break
            else:
                print(" Invalid choice! Please select a valid number.")
        except ValueError:
            print(" Invalid input! Please enter a number.")

    
        file_path = input("\nEnter the full path of the .docx file: ")
        print(f" Debug: File path received -> {file_path}")

    
        translated_file = translate_and_save(file_path, target_language)
        
        return jsonify({'downloadUrl': f"http://localhost:5000/download/{translated_path}"})

@app.route('/download/<path:filename>', methods=['GET'])
def download_file(filename):
    return app.send_static_file(filename)

if __name__ == '__main__':
    app.run(debug=True, port=8000)